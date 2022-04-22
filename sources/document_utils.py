#!/bin/python3

from docx import Document
from docx.shared import Inches, Cm
from docx.table import Table

DOCUMENT = Document()


def center() -> None:
    """
    It centers the last paragraph in the document
    """
    DOCUMENT.paragraphs[-1].alignment = 1


def justify() -> None:
    DOCUMENT.paragraphs[-1].alignment = 3


def add_picture(path: str, width: float = None, height: float = None) -> None:
    DOCUMENT.add_picture(path, width=Inches(width), height=Inches(height))


def add_picture_center(path: str, width: float = None, height: float = None) -> None:
    add_picture(path, width, height)
    center()


def add_title(title: str, level: int = 1) -> None:
    DOCUMENT.add_heading(title, level=level)


def add_title_center(title: str, level: int = 1) -> None:
    add_title(title, level)
    center()


def add_paragraph(text: str) -> None:
    DOCUMENT.add_paragraph(text)
    justify()


def add_paragraph_indent(text: str, indent: int = 1) -> None:
    add_paragraph(text)
    DOCUMENT.paragraphs[-1].paragraph_format.first_line_indent = Cm(indent)


def create_array(rows: int, cols: int) -> Table:
    return DOCUMENT.add_table(rows=rows, cols=cols)


def page_break() -> None:
    DOCUMENT.add_page_break()


def save(path: str):
    DOCUMENT.save(path)
